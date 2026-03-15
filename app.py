import streamlit as st
import boto3
from botocore.exceptions import NoCredentialsError
from PIL import Image
import io
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

MODEL_ID = "amazon.nova-lite-v1:0"

st.set_page_config(page_title="AI Screenshot Study Explainer", layout="wide")

st.title("📚 AI Screenshot Study Explainer")

st.sidebar.header("Settings")

difficulty = st.sidebar.selectbox(
    "Choose explanation difficulty",
    ["Beginner", "Intermediate", "Advanced"]
)

uploaded_file = st.file_uploader("Upload a study screenshot", type=["png","jpg","jpeg"])

result_text = ""

if uploaded_file:

    image = Image.open(uploaded_file)
    st.image(image, caption="Uploaded Screenshot", use_column_width=True)

    if st.button("Explain Screenshot"):

        try:
            client = boto3.client("bedrock-runtime", region_name="us-east-1")

            img_bytes = io.BytesIO()
            image.save(img_bytes, format="PNG")
            img_bytes = img_bytes.getvalue()

            prompt = f"""
            Explain the content of this study screenshot in {difficulty} level.
            Also generate:
            1. Key study notes
            2. 3 quiz questions
            """

            response = client.converse(
                modelId=MODEL_ID,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"text": prompt},
                            {"image": {"format": "png", "source": {"bytes": img_bytes}}},
                        ],
                    }
                ],
            )

            result_text = response["output"]["message"]["content"][0]["text"]

        except NoCredentialsError:
            result_text = "⚠️ AWS credentials not configured yet. Please run 'aws configure'."

        except Exception as e:
            result_text = f"Error: {str(e)}"

    if result_text:
        st.subheader("Explanation")
        st.write(result_text)

        # ---------- PDF GENERATION ----------
        def generate_pdf(text):
            buffer = io.BytesIO()
            styles = getSampleStyleSheet()

            story = []
            for line in text.split("\n"):
                story.append(Paragraph(line, styles["Normal"]))
                story.append(Spacer(1,10))

            doc = SimpleDocTemplate(buffer)
            doc.build(story)

            buffer.seek(0)
            return buffer

        pdf_file = generate_pdf(result_text)

        st.download_button(
            label="⬇ Download as PDF",
            data=pdf_file,
            file_name="study_explanation.pdf",
            mime="application/pdf"
        )

        # ---------- DOCX GENERATION ----------
        def generate_doc(text):
            doc = Document()
            doc.add_heading("Study Explanation", level=1)

            for line in text.split("\n"):
                doc.add_paragraph(line)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        doc_file = generate_doc(result_text)

        st.download_button(
            label="⬇ Download as DOCX",
            data=doc_file,
            file_name="study_explanation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
